from json import JSONDecodeError, loads, load, dumps, dump
from os import environ, path, makedirs

from dotenv import load_dotenv

from loguru import logger

import openai

from docx import Document
import tiktoken

load_dotenv()


class DataCategorizer:
    openai.api_key = environ.get('OPENAI_API_KEY')

    def count_tokens(self, text, model="gpt-3.5-turbo"):
        """Estimate the number of tokens in a text using tiktoken."""
        encoding = tiktoken.get_encoding("cl100k_base")  # Encoding for GPT-3.5 turbo
        tokens = encoding.encode(text)
        return len(tokens)

    def split_text_into_chunks(self, text, max_tokens=4096):
        """Split text into chunks that fit within the token limits."""
        chunks = []
        current_chunk = []
        current_chunk_tokens = 0

        sentences = text.split('. ')

        for sentence in sentences:
            tokens = self.count_tokens(sentence)
            if current_chunk_tokens + tokens <= max_tokens:
                current_chunk.append(sentence)
                current_chunk_tokens += tokens
            else:
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_chunk_tokens = tokens
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        return chunks

    async def corruption_categorizer(self, text):
        """Determine if the article is related to corruption."""
        prompt = (
            f'Ви отримаєте текст новини українською мовою: {text}. '
            'Вам потрібно визначити, чи є цей текст пов\'язаний з корупцією. '
            'Якщо текст пов\'язаний з корупцією, відповідь повинна бути "True". '
            'Якщо текст не пов\'язаний з корупцією, відповідь повинна бути "False".'
        )

        chunks = self.split_text_into_chunks(prompt)

        results = []
        for chunk in chunks:
            response = openai.ChatCompletion.create(
                model='gpt-3.5-turbo',
                messages=[
                    {'role': 'system', 'content': 'Ви корисний асистент, який визначає корупційний контент.'},
                    {'role': 'user', 'content': chunk},
                ],
                max_tokens=150,
                temperature=0.5,
                timeout=6000
            )
            result = response['choices'][0]['message']['content'].strip()
            results.append(result)

        return "True" if "True" in results else "False"

    async def corruption_data_only(self, text):
        """Extract detailed corruption-related data from the article."""
        prompt = (
            f"Ви отримаєте текст новини українською мовою: {text}. "
            "Будь ласка, екстрагуйте та надайте наступну інформацію у форматі JSON:\n"
            "1. **Індивідууми**: Перерахуйте всіх осіб, які згадуються, їхні посади та афіліації (якщо є).\n"
            "2. **Юридичні особи**: Перерахуйте всі компанії або організації в Україні (приватні або державні).\n"
            "3. **Офшори**: Якщо згадуються офшорні компанії або активи, перерахуйте їх.\n"
            "4. **Державні органи**: Перерахуйте будь-які залучені державні органи.\n\n"
            "Формат вихідних даних:\n"
            "{\n"
            "  \"individuals\": [\n"
            "    { \"name\": \"Ім'я особи\", \"position\": \"Посада\", \"affiliations\": [ \"Організація\" ] }\n"
            "  ],\n"
            "  \"legal_entities\": [\n"
            "    { \"entity\": \"Назва організації\", \"type\": \"Державна або приватна\" }\n"
            "  ],\n"
            "  \"offshore\": [ \"Офшорна компанія\" ],\n"
            "  \"government_bodies\": [ \"Державний орган\" ]\n"
            "}"
        )

        chunks = self.split_text_into_chunks(prompt)

        results: list[dict[str, str]] = []
        for chunk in chunks:
            response = openai.ChatCompletion.create(
                model='gpt-3.5-turbo',
                messages=[
                    {'role': 'system', 'content': 'Ви корисний асистент, який визначає корупційний контент.'},
                    {'role': 'user', 'content': chunk},
                ],
                max_tokens=500,
                temperature=0.5,
                timeout=6000
            )
            result = response['choices'][0]['message']['content'].strip()
            try:
                corruption_data = loads(result)
                results.append(corruption_data)
            except JSONDecodeError:
                results.append({"message": "No corruption-related data found."})
        logger.info(results)
        return results

    async def corruption_schemes(self, file_data):
        prompt = (
            f"Ти отримуєш json текст, пов'язаний із різного виду корупцією: {file_data}. "
            "Будь ласка, витягни з наданого тексту схеми корупції, класифікуючи їх за типами, які наведені в описі. "
            "Ось приклади можливих типів корупції з описами. Залишіть тільки ті схеми, які точно згадані в тексті. "
            "Ось види корупції та приклади можливих корупційних схем:\n"
            "1) **Корупція в сфері оборони**: фіктивні тендери Міноборони, корупція в закупівлях для ЗСУ, фіктивні контракти Міноборони, неякісна техніка для ЗСУ, контрабанда комплектуючих для ЗСУ, непрозорі оборонні контракти, тіньові схеми постачання зброї тощо.\n"
            "2) **Контрабанда**: схеми на митниці, відкат на митниці, зникнення вантажів на митниці, офшорні схеми імпорту.\n"
            "3) **Зловживання в державних закупівлях**: тендерні махінації, відкати на держзакупівлях, зловживання при закупівлях, тендерні змови, завищення цін при держзакупівлі.\n"
            "4) **Незаконна приватизація**: дерибан (або ж розкрадання) державного майна, маніпуляції при оцінці державного майна, заниження вартості об’єктів.\n"
            "5) **Розкрадання кредитів державних банків**: розкрадання кредитів, виведення кредитних коштів, провалені кредитні програми, фіктивні кредити.\n"
            "6) **АРМА та державне рейдерство**: виведення активів через АРМА, заниження вартості активів, державне рейдерство.\n"
            "7) **Антимонопольний комітет України (далі: АМКУ) та перерозподіл ринків**: політичний вплив на керівництво АМКУ, лобіювання інтересів окремих фінансово-промислових груп, формальний характер перевірок зловживань монопольним становищем.\n"
            "8) **Розкрадання державного майна**: низька прозорість процесів інвентаризації та передачі державного майна, системна корупція серед посадових осіб, відповідальних за облік і збереження майна, виведення держмайна за кордон.\n"
            "9) **Незаконний видобуток природних ресурсів**: незаконний видобуток та контрабанда природних ресурсів (бурштин, нафти, газу).\n"
            "10) **Зловживання службовим становищем**: корупція посадовців, виведення коштів через службові рішення, лобізм та зловживання.\n"
            "11) **Зловживання при розподілі земельних ресурсів**: прихована приватизація землі, корупція та обхід на земельних аукціонах, виведення сільгоспземель під забудову.\n"
            "12) **Корупція в містобудуванні**: корупція в будівництві, відкати при узгодженні проектів, офшори на будівництві, незаконне будівництво.\n"
            "13) **Корупція в правоохоронних органах**: фальсифікація справ, хабарі слідчим, правоохоронна мафія, маніпуляції з доказами, корупція в ДБР.\n"
            "14) **Корупція в судах**: відкати за рішення, легалізація рішень за хабарі, зловживання суддівськими повноваженнями.\n"
            "15) **Розкрадання гуманітарної та/або військової допомоги**: крадіжка гуманітарної допомоги, маніпулювання наданням допомоги для власної вигоди, продаж на чорному ринку."
        )
        chunks = self.split_text_into_chunks(prompt)
        results = []

        for chunk in chunks:
            try:
                response = openai.ChatCompletion.create(
                    model='gpt-3.5-turbo',
                    messages=[
                        {'role': 'system', 'content': 'Ви корисний асистент, який визначає корупційний контент.'},
                        {'role': 'user', 'content': chunk},
                    ],
                    max_tokens=500,
                    temperature=0.5,
                    timeout=6000
                )
                result = response['choices'][0]['message']['content'].strip()
                corruption_schemes = []
                for line in result.split("\n"):
                    if line.strip():
                        parts = line.split(":", 1)
                        if len(parts) == 2 and parts[1].strip():
                            corruption_type = parts[0].split(")", 1)[
                                1].strip()
                            corruption_schemes.append({
                                "type": corruption_type,
                                "message": parts[1].strip()
                            })
                if corruption_schemes:
                    results.append({"corruption_schemes": corruption_schemes})
            except Exception as e:
                results.append({"error": f"Error processing chunk: {str(e)}"})
        return results

    async def process_json_file(self, input_file_path, output_file_path_json, output_file_path_docx):
        with open(input_file_path, 'r', encoding='utf-8') as file:
            data = load(file)

        if isinstance(data, list) and len(data) > 0:
            results = []
            doc = Document()
            doc.add_heading('Corruption Data Report', 0)

            for article in data:
                text = article.get("short_text", "")
                link = article.get("link", "Не надано")
                logger.info(f"Extracted link: {link}")

                article_info = {
                    "date": article.get("date", "Не надано"),
                    "link": link,
                    "title": article.get("title", "Не надано"),
                    "author": article.get("author", "Не надано"),
                    "short_text": text
                }

                # Step 1: Check if the article is related to corruption by checking the entire article text
                corruption_schemes = await self.corruption_schemes(text)

                if corruption_schemes:  # Only proceed if corruption schemes are found
                    # Flatten the nested corruption_schemes to avoid redundancy
                    flat_corruption_schemes = []
                    seen_types = set()  # To track types and avoid duplicates

                    for scheme_group in corruption_schemes:
                        for scheme in scheme_group.get("corruption_schemes", []):
                            if scheme["type"] not in seen_types:
                                flat_corruption_schemes.append(scheme)
                                seen_types.add(scheme["type"])

                    # Add the unique corruption schemes to the article info
                    article_info["corruption_schemes"] = flat_corruption_schemes

                    # Step 2: Split the article text into chunks if it's too long
                    chunks = self.split_text_into_chunks(text)

                    combined_extracted_data = []
                    for chunk in chunks:
                        # Extract the corruption-related data from the chunk
                        extracted_data = await self.corruption_data_only(chunk)
                        if not extracted_data:
                            combined_extracted_data.append({"message": "No corruption-related data found."})
                        else:
                            combined_extracted_data.extend(extracted_data)

                    # Add extracted corruption data to the article
                    article_info["corruption_data"] = combined_extracted_data

                    # Append the article with corruption data to the results
                    results.append(article_info)

                    # Add the article data to the DOCX file
                    doc.add_heading(article_info["title"], level=1)
                    doc.add_paragraph(f"Date: {article_info['date']}")
                    doc.add_paragraph(f"Link: {article_info['link']}")
                    doc.add_paragraph(f"Author: {article_info['author']}")
                    doc.add_paragraph(f"Short Text: {article_info['short_text']}")

                    # Add corruption schemes in the DOCX file
                    for scheme in flat_corruption_schemes:
                        doc.add_paragraph(f"Corruption Type: {scheme['type']}")
                        doc.add_paragraph(f"Message: {scheme['message']}")

            # Step 3: Write the results to the output JSON file
            with open(output_file_path_json, 'w', encoding='utf-8') as json_file:
                dump(results, json_file, ensure_ascii=False, indent=4)

            # Step 4: Save the DOCX file
            doc.save(output_file_path_docx)
